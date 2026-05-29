#!/usr/bin/env python3
"""Simple Markdown to DOCX converter for Ольга's reports."""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_runs_with_formatting(paragraph, text):
    """Parse inline formatting: **bold**, *italic*, [link](url), `code`."""
    # Pattern matches bold, italic, links, code in order of priority
    pattern = re.compile(
        r'\*\*([^\*]+)\*\*'          # **bold**
        r'|\*([^\*]+)\*'             # *italic*
        r'|`([^`]+)`'                # `code`
        r'|\[([^\]]+)\]\(([^\)]+)\)' # [text](url)
    )
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group(1) is not None:  # bold
            run = paragraph.add_run(m.group(1))
            run.bold = True
        elif m.group(2) is not None:  # italic
            run = paragraph.add_run(m.group(2))
            run.italic = True
        elif m.group(3) is not None:  # code
            run = paragraph.add_run(m.group(3))
            run.font.name = 'Courier New'
        elif m.group(4) is not None:  # link
            run = paragraph.add_run(m.group(4))
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            run.underline = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def parse_table(lines, start_idx):
    """Parse markdown table starting at start_idx. Returns (rows, end_idx)."""
    rows = []
    i = start_idx
    while i < len(lines) and '|' in lines[i]:
        line = lines[i].strip()
        if re.match(r'^\|?[\s\-:|]+\|?$', line):  # separator row
            i += 1
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
        i += 1
    return rows, i


def md_to_docx(md_path, docx_path):
    doc = Document()
    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    i = 0
    in_code = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        # Code blocks
        if stripped.startswith('```'):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Headers
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith('#### '):
            doc.add_heading(stripped[5:].strip(), level=4)
        # Horizontal rule
        elif stripped in ('---', '***', '___'):
            p = doc.add_paragraph()
            p.add_run('─' * 60).font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        # Tables
        elif '|' in stripped and stripped.strip().startswith('|'):
            rows, new_i = parse_table(lines, i)
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Light Grid Accent 1'
                for r_idx, row in enumerate(rows):
                    for c_idx, cell_text in enumerate(row):
                        if c_idx < len(table.rows[r_idx].cells):
                            cell = table.rows[r_idx].cells[c_idx]
                            cell.text = ''
                            p = cell.paragraphs[0]
                            add_runs_with_formatting(p, cell_text)
                            if r_idx == 0:
                                for run in p.runs:
                                    run.bold = True
                i = new_i
                continue
        # Blockquote
        elif stripped.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.75)
            text = stripped[2:]
            run = p.add_run('▎ ')
            run.font.color.rgb = RGBColor(0x4A, 0x90, 0xE2)
            run.bold = True
            add_runs_with_formatting(p, text)
            for run in p.runs[1:]:
                run.italic = True
        # Ordered list
        elif re.match(r'^\s*\d+\.\s', line):
            indent_level = (len(line) - len(line.lstrip())) // 2
            text = re.sub(r'^\s*\d+\.\s', '', line)
            p = doc.add_paragraph(style='List Number')
            if indent_level > 0:
                p.paragraph_format.left_indent = Cm(0.75 * indent_level)
            add_runs_with_formatting(p, text)
        # Unordered list
        elif re.match(r'^\s*[-*]\s', line):
            indent_level = (len(line) - len(line.lstrip())) // 2
            text = re.sub(r'^\s*[-*]\s', '', line)
            p = doc.add_paragraph(style='List Bullet')
            if indent_level > 0:
                p.paragraph_format.left_indent = Cm(0.75 * (indent_level + 1))
            add_runs_with_formatting(p, text)
        # Empty line
        elif not stripped:
            pass
        # Regular paragraph
        else:
            p = doc.add_paragraph()
            add_runs_with_formatting(p, stripped)

        i += 1

    doc.save(docx_path)
    print(f'✓ {docx_path}')


if __name__ == '__main__':
    for md_file in sys.argv[1:]:
        md_path = Path(md_file)
        docx_path = md_path.with_suffix('.docx')
        md_to_docx(md_path, docx_path)
